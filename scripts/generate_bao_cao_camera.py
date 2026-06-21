from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
REPORTS = BASE / "reports"
OUT = REPORTS / "Bao_cao_BTL_Camera_Stream_Service.docx"


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.05)
    section.right_margin = Inches(0.8)

    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.2
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 20, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 14, "2F5597"),
        ("Heading 3", 13, "2F5597"),
    ]:
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def pb(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def p(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.add_run(text)


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, headers, rows) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(tbl.rows[0].cells[i], header, True)
        tc_pr = tbl.rows[0].cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9EAF7")
        tc_pr.append(shd)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph()


def image(doc: Document, filename: str, caption: str) -> None:
    path = REPORTS / filename
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(5.8))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True


CORE_SECTIONS = [
    (
        "CHUONG 1. TONG QUAN VA CO SO NGHIEP VU",
        [
            ("1.1. Boi canh Smart Campus Operations Platform", "Smart Campus Operations Platform la bai toan mo phong mot he thong van hanh truong hoc thong minh, trong do moi nhom phu trach mot service rieng. Cac service khong lam viec doc lap hoan toan ma phai trao doi du lieu voi nhau theo dependency map. Camera Stream lien quan truc tiep den AI Vision, Analytics va Core Business. Diem quan trong cua bai tap lon khong chi la viet API chay tren may ca nhan, ma la chung minh service co hop dong ro rang, co bang chung kiem thu va co kha nang tich hop voi nhom khac."),
            ("1.2. Vi tri Camera Stream trong dependency map", "Camera Stream la service nam giua nguon camera va cac service phan tich. Lien ket Camera Stream -> AI Vision dung REST sync vi Camera can ket qua detect ngay de tao event da phan tich. Lien ket Camera Stream -> Analytics dung queue async/MQTT vi Analytics chi can aggregate event ve sau. Voi Core Business, Camera Stream chi nen cung cap tin hieu risk/unknown_person/alert_candidate, khong xu ly policy cuoi cung thay Core."),
            ("1.3. Bai toan nghiep vu cua service camera", "Bai toan nghiep vu la bien mot luong camera lien tuc thanh nhung frame va event co y nghia. Service khong gui moi frame sang AI, ma chi xu ly khi co motion, gan metadata, encode anh base64 va goi Vision. Sau khi Vision tra detections, unknown_person va risk_level, Camera Stream tao event camera.motion.analyzed de nhom Analytics va Core Business co the su dung."),
            ("1.4. Du lieu dau vao va dau ra", "Dau vao cua Camera Stream gom frame camera va metadata toi thieu nhu camera_id, location, captured_at, motion_score. Dau ra gom response API cho nguoi goi va event lien service cho Analytics. Response API phuc vu demo va kiem thu truc tiep; event lien service phuc vu pipeline bat dong bo. Hai loai du lieu nay can duoc tach ro de tranh viec gui anh base64 sang nhung noi chi can aggregate."),
            ("1.5. Ranh gioi voi cac service khac", "Camera Stream khong lam thay AI Vision, khong tinh KPI thay Analytics va khong xu ly policy an ninh thay Core Business. Trong lab, Vision la mock de kiem thu hop dong, khong phai AI model that. Analytics nhan event da xu ly de tinh KPI, khong nhan raw image. Core Business moi la noi ket hop camera event voi Access Gate va policy ra vao."),
            ("1.6. Muc tieu va pham vi bao cao", "Bao cao tong hop phan nghiep vu, thiet ke, contract, demo va kiem thu cua Camera Stream Service. Pham vi bao gom OpenAPI, API upload/analyze frame, Docker Compose, Newman evidence, script chup frame live, script publish MQTT sang Analytics va huong dan du lieu gui cho nhom Vision/Analytics. Ngoai pham vi la AI model that va worker frame-difference production day du."),
        ],
    ),
    (
        "CHUONG 2. PHAN TICH THIET KE VA CAI DAT CAMERA STREAM",
        [
            ("2.1. Kien truc tong the", "Kien truc thuc hanh gom Camera API, PostgreSQL, AI Vision mock va Analytics mock trong Docker Compose. API tiep nhan request, database luu metadata frame, Vision mock gia lap detect va Analytics mock gia lap noi nhan event. Folder nen dung de demo la C:\\Projects\\Bailap-dich_Vu_Ket_Noi\\lab-5-minhtan39 vi duong dan nay tranh loi Docker voi OneDrive co dau tieng Viet."),
            ("2.2. Quy trinh phat hien chuyen dong va chup frame", "Quy trinh bat dau tu camera source. Service/script doc frame, kiem tra frame hop le, gan camera_id, location va captured_at, sau do tinh hoac gan motion_score. Khi motion_score vuot nguong, anh duoc nen thanh JPEG/PNG, encode base64 va gui sang Vision. Phien ban demo da co script chup frame live; phien ban production nen bo sung worker doc stream lien tuc va so sanh frame truoc/sau."),
            ("2.3. Hop dong API noi bo", "OpenAPI 3.1.0 dinh nghia cac endpoint /health, /api/v1/frames, /api/v1/frames/{frame_id} va /api/v1/frames/{frame_id}/analyze. Schema FrameCreate yeu cau camera_id dung mau CAM-A01, frame_format la jpeg/png, image_base64 co gioi han kich thuoc va motion_score nam trong khoang 0 den 1. Neu request sai, service tra Problem Details thay vi loi chung chung."),
            ("2.4. Hop dong Camera -> AI Vision", "Camera Stream goi Vision qua POST /api/v1/detect. Payload gom request_id, camera_id, timestamp, location, motion_score, image_base64 va snapshot_url tuy chon. Trong lab, nhom dung image_base64 vi cac nhom khong co file hosting chung. Response cua Vision gom detections, unknown_person va risk_level. Camera Stream luu ket qua nay de tao event da phan tich."),
            ("2.5. Event Camera -> Analytics", "Sau khi analyze, Camera Stream tao event camera.motion.analyzed. Event gom source_service, request_id, frame_id, camera_id, location, occurred_at, motion_detected, motion_score, motion_level, risk_level, unknown_person va alert_candidate. Event nay duoc publish async/MQTT sang topic smart-campus/events/camera de Analytics aggregate KPI."),
            ("2.6. Xu ly loi va Problem Details", "Service phai xu ly loi ro rang: 401 khi thieu token, 422 khi validation sai, 404 khi frame_id khong ton tai, 502 khi dependency tra loi loi va 503 khi dependency timeout/unavailable. Dinh dang application/problem+json giup consumer hieu loi va giup Newman/GitHub Actions kiem tra on dinh."),
            ("2.7. Bao mat va kiem soat truy cap", "Du lieu camera co tinh nhay cam nen service can bearerAuth, gioi han kich thuoc image_base64, khong log toan bo anh, rate limit/throttle theo camera_id va chi chia se event can thiet sang Analytics. Neu sau nay dung snapshot_url, URL can co quyen truy cap va thoi gian het han."),
            ("2.8. Cau truc du an va Docker Compose", "Du an gom src/camera_app, contracts, reports, scripts va docker-compose.yml. Compose stack gom fit4110-camera-api-lab05, fit4110-camera-db-lab05, fit4110-camera-vision-mock-lab05 va fit4110-camera-analytics-mock-lab05. Bang chung trong reports gom docker-compose-ps.png, health-api.png, camera-live-source.png, analyze-live-frame-response.png va Newman HTML/XML."),
            ("2.9. Cac script demo tu dong", "auto_capture_camera.py lay frame live, ghi file auto-camera-frame.jpg, encode base64, upload vao API va goi analyze. publish_camera_event_demo.py gui event mau sang MQTT Analytics. auto_capture_camera_mqtt_analytics.py ket hop luong day du: chup frame, upload API, analyze qua Vision mock va publish event sang MQTT. Day la script nen dung khi can demo lien nhom."),
            ("2.10. Dieu chinh so voi nghiep vu lien nhom", "Theo trao doi lien nhom, du lieu gui cho Vision phai co anh base64 kem metadata; du lieu gui cho Analytics phai la event da xu ly va khong chua anh. Cach chia nay giup Vision lam detect, Analytics lam aggregate, Core Business lam policy. Day cung la diem can nhan manh khi lam Word va thuyet trinh."),
        ],
    ),
    (
        "CHUONG 3. DEMO, KIEM THU VA DANH GIA KET QUA",
        [
            ("3.1. Moi truong demo", "Moi truong demo chay tai C:\\Projects\\Bailap-dich_Vu_Ket_Noi\\lab-5-minhtan39. Can bat Docker Desktop, chay docker compose up -d --build, sau do dung docker compose ps de kiem tra API, DB, Vision mock va Analytics mock. Endpoint kiem tra nhanh la http://localhost:8000/health, http://localhost:9000/health va http://localhost:9010/health."),
            ("3.2. Kich ban demo chinh", "Kich ban demo gom bon buoc: khoi dong stack, kiem tra health, chay script chup frame live/upload/analyze, va publish event sang Analytics. Khi thuyet trinh, can noi theo luong nghiep vu: co motion -> chup frame -> encode base64 -> gui Vision -> nhan detections -> tao event Analytics."),
            ("3.3. Du lieu gui cho nhom Vision", "Du lieu gui cho Vision la detect request co image_base64 va context. image_base64 la anh, con request_id, camera_id, timestamp, location va motion_score giup Vision log va doi soat. Trong demo, Vision mock tra detections, unknown_person va risk_level; Camera Stream khong tu nhan dien AI ma chi dung ket qua de tao event."),
            ("3.4. Du lieu gui cho nhom Analytics", "Khi nhom Analytics hoi 'gui du lieu nhom ban da xu ly', du lieu dung can gui la event camera.motion.analyzed, khong phai anh raw. Event khong chua image_base64 de payload nhe va dung muc dich aggregate. Cac truong quan trong la camera_id, location, motion_score, motion_level, risk_level, unknown_person va alert_candidate."),
            ("3.5. Ket qua kiem thu API", "Bang chung Newman cho thay collection lab05 compose da chay 8 request, 15 assertions va 0 failed trong lan ghi nhan. Cac test bao gom health, upload frame, validation, analyze qua Vision mock va tao analytics_event. Bang chung anh chup trong reports giup chung minh stack va endpoint chay duoc."),
            ("3.6. Bang doi chieu yeu cau", "Yeu cau REST sync Camera -> Vision da dap ung qua endpoint analyze. Yeu cau anh base64 da dap ung qua VisionDetectRequest. Yeu cau async Camera -> Analytics da demo qua event-shaped payload va script MQTT. Yeu cau Docker/Newman/evidence da co trong reports. Diem chua production la AI model that va worker motion lien tuc."),
            ("3.7. Han che hien tai", "Han che lon nhat la Vision dang la mock, motion detection lien tuc chua production-ready va publish async chua co durable retry/dead-letter queue. Cac han che nay khong lam hong muc tieu lab, nhung can ghi ro de bao cao trung thuc va de thay biet nhom nam duoc ranh gioi giua demo va he thong that."),
            ("3.8. Huong phat trien", "Huong phat trien la bo sung worker frame-difference lien tuc, cooldown theo camera_id, retry/circuit breaker khi Vision loi, schema version cho event MQTT, message durability va dashboard Analytics. Neu Vision that san sang, Camera Stream chi can doi dependency URL ma giu nguyen contract detect request."),
        ],
    ),
]


def add_expand(doc: Document, title: str) -> None:
    extras = [
        "Ve mat nghiep vu, phan nay giup chung minh service khong chi co code ma con co ly do ton tai trong toan bo he thong. Moi truong lien nhom yeu cau moi truong du lieu ro rang, vi neu mot field bi thieu thi service tiep theo se khong the xu ly dung.",
        "Ve mat thiet ke, cach tach metadata, anh base64 va event giup Camera Stream giu dung boundary. Anh base64 chi phuc vu Vision, con Analytics nhan event rut gon. Day la diem quan trong de tranh nham lan voi service khac trong bai lien nhom.",
        "Ve mat demo, noi dung nay can duoc gan voi bang chung trong thu muc reports. Khi thuyet trinh, nen mo terminal va anh evidence theo dung thu tu: health, upload frame, analyze response, publish event. Cach trinh bay nay giup thay de theo doi luong nghiep vu.",
    ]
    for item in extras:
        p(doc, item)


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold in [
        ("TRUONG DAI HOC / KHOA CONG NGHE THONG TIN", 13, True),
        ("HOC PHAN FIT4110 - CONNECTIVITY SERVICES", 13, True),
        ("", 12, False),
        ("BAO CAO BAI TAP LON", 20, True),
        ("SMART CAMPUS OPERATIONS PLATFORM", 18, True),
        ("SERVICE CAMERA STREAM", 18, True),
    ]:
        run = cover.add_run(text + "\n")
        run.bold = bold
        run.font.size = Pt(size)
    for line in [
        "Nhom phu trach: Camera Stream Service",
        "Nguoi nop/tai khoan: minhtan39",
        "Noi dung: phat hien chuyen dong, chup frame, gui anh base64 sang AI Vision, publish event sang Analytics",
        "Ngay tao bao cao: " + datetime.now().strftime("%d/%m/%Y %H:%M"),
    ]:
        para = doc.add_paragraph(line)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pb(doc)

    h(doc, "LOI NOI DAU")
    for text in [
        "Bao cao nay tong hop phan bai tap lon cua nhom Camera Stream trong Smart Campus Operations Platform. Trong bo nghiep vu 7 service, Camera Stream dong vai tro thu nhan du lieu hinh anh tu camera, bien luong camera thanh cac frame co y nghia, va chuyen frame do sang AI Vision de phan tich.",
        "Trong qua trinh lam lab, nhom da bo sung phan chup chuyen dong va tra anh dang base64 de gui cho team AI Vision. Sau khi co ket qua detect, service tao event camera.motion.analyzed de Analytics co the aggregate KPI va Core Business co the ap dung policy neu can.",
        "Bao cao duoc viet theo cau truc 3 chuong: tong quan nghiep vu, phan tich thiet ke/cai dat, va demo/kiem thu/danh gia. Cac phu luc cuoi tai lieu cung cap lenh demo, payload mau va ma tran test de dung khi nop bai hoac thuyet trinh.",
    ]:
        p(doc, text)
    pb(doc)

    h(doc, "MUC LUC")
    for line in [
        "LOI NOI DAU",
        "DANH MUC TU VIET TAT",
        "DANH MUC BANG VA HINH",
        "MO DAU",
        "CHUONG 1. TONG QUAN VA CO SO NGHIEP VU",
        "CHUONG 2. PHAN TICH THIET KE VA CAI DAT CAMERA STREAM",
        "CHUONG 3. DEMO, KIEM THU VA DANH GIA KET QUA",
        "KET LUAN",
        "TAI LIEU THAM KHAO",
        "PHU LUC",
    ]:
        doc.add_paragraph(line)
    pb(doc)

    h(doc, "DANH MUC TU VIET TAT")
    table(
        doc,
        ["Tu viet tat", "Y nghia", "Giai thich trong de tai"],
        [
            ["API", "Application Programming Interface", "Giao dien de cac service goi nhau"],
            ["REST", "Representational State Transfer", "Co che goi dong bo Camera -> Vision"],
            ["MQTT", "Message Queuing Telemetry Transport", "Co che publish event bat dong bo sang Analytics"],
            ["Base64", "Binary-to-text encoding", "Cach ma hoa anh de nhung truc tiep trong JSON"],
            ["AI Vision", "Dich vu phan tich hinh anh", "Provider nhan frame tu Camera Stream"],
            ["Mock", "Dich vu gia lap", "Vision/Analytics trong lab dung mock de test tich hop"],
        ],
    )
    pb(doc)

    h(doc, "DANH MUC BANG VA HINH")
    table(
        doc,
        ["Ma", "Ten bang/hinh", "Noi dung"],
        [
            ["Bang 1", "Vai tro cac service lien quan", "Phan dinh Camera, Vision, Analytics, Core"],
            ["Bang 2", "Endpoint API Camera Stream", "Cac endpoint health, frame, analyze"],
            ["Bang 3", "Payload Vision", "Cau truc request Camera gui sang Vision"],
            ["Bang 4", "Event Analytics", "Cau truc event camera.motion.analyzed"],
            ["Bang 5", "Ma tran kiem thu", "Kich ban Newman va demo thu cong"],
            ["Hinh 1", "Docker Compose status", "Bang chung cac container chay healthy"],
            ["Hinh 2", "Health API", "Bang chung endpoint /health"],
            ["Hinh 3", "Frame camera live", "Frame demo tu camera source"],
            ["Hinh 4", "Analyze response", "Ket qua phan tich frame"],
        ],
    )
    pb(doc)

    h(doc, "MO DAU")
    for text in [
        "Trong cac nen tang campus thong minh, camera khong chi la thiet bi ghi hinh ma con la nguon du lieu van hanh quan trong. Neu xu ly dung, camera co the ho tro phat hien chuyen dong bat thuong, ghi nhan su kien an ninh va cung cap du lieu cho he thong thong ke.",
        "Bai tap lon tap trung vao cach cac service doc lap phoi hop voi nhau thong qua hop dong API va event. Doi voi nhom Camera Stream, yeu cau quan trong la mo ta ro provider/consumer, endpoint trao doi, du lieu dau vao dau ra, cach xu ly loi va bang chung demo.",
        "Pham vi bao cao khong trinh bay Word nhu mot phan source code, ma tong hop tat ca noi dung can dung cho nop bai: phan tich nghiep vu, thiet ke ky thuat, contract, demo, kiem thu, bang chung anh chup, han che va huong phat trien.",
    ]:
        p(doc, text)
    pb(doc)

    for chapter, sections in CORE_SECTIONS:
        h(doc, chapter)
        pb(doc)
        for title, body in sections:
            h(doc, title, 2)
            p(doc, body)
            add_expand(doc, title)
            if title == "1.2. Vi tri Camera Stream trong dependency map":
                table(
                    doc,
                    ["Consumer", "Provider", "Muc dich", "Co che"],
                    [
                        ["Camera Stream", "AI Vision", "Gui frame khi phat hien motion", "REST sync"],
                        ["Camera Stream", "Analytics", "Feed event camera cho aggregate", "Queue async/MQTT"],
                        ["Core Business", "Camera Stream", "Nhan tin hieu risk/policy khi can", "REST/event"],
                    ],
                )
            if title == "2.3. Hop dong API noi bo":
                table(
                    doc,
                    ["Endpoint", "Method", "Chuc nang", "Ket qua chinh"],
                    [
                        ["/health", "GET", "Kiem tra readiness", "200 status ok"],
                        ["/api/v1/frames", "POST", "Upload frame base64", "201 FrameAccepted"],
                        ["/api/v1/frames", "GET", "Xem danh sach frame", "200 FramePage"],
                        ["/api/v1/frames/{frame_id}", "GET", "Xem chi tiet frame", "200 hoac 404"],
                        ["/api/v1/frames/{frame_id}/analyze", "POST", "Goi Vision va tao event", "200 AnalyzeResult"],
                    ],
                )
            if title == "2.4. Hop dong Camera -> AI Vision":
                code(
                    doc,
                    'POST /api/v1/detect\n{\n  "request_id": "vision-FR-20260620-0001",\n  "camera_id": "CAM-A01",\n  "timestamp": "2026-06-20T08:30:00+07:00",\n  "location": "Main lobby",\n  "motion_score": 0.82,\n  "image_base64": "...",\n  "snapshot_url": null\n}',
                )
            if title == "2.5. Event Camera -> Analytics":
                code(
                    doc,
                    'Topic: smart-campus/events/camera\n{\n  "event_type": "camera.motion.analyzed",\n  "source_service": "team-camera",\n  "frame_id": "FR-20260620-0001",\n  "camera_id": "CAM-A01",\n  "motion_score": 0.82,\n  "motion_level": "high",\n  "risk_level": "high",\n  "unknown_person": true,\n  "alert_candidate": true\n}',
                )
            if title == "3.1. Moi truong demo":
                image(doc, "docker-compose-ps.png", "Hinh 1. Trang thai Docker Compose cua cac container demo")
                image(doc, "health-api.png", "Hinh 2. Ket qua kiem tra endpoint health cua Camera API")
            if title == "3.3. Du lieu gui cho nhom Vision":
                image(doc, "camera-live-source.png", "Hinh 3. Frame camera live dung trong demo")
                image(doc, "analyze-live-frame-response.png", "Hinh 4. Response analyze sau khi gui frame sang Vision mock")
            if title == "3.5. Ket qua kiem thu API":
                table(
                    doc,
                    ["Nhom test", "Noi dung", "Ket qua ghi nhan"],
                    [
                        ["Health", "API va dependency health endpoint", "Pass"],
                        ["Frame upload", "POST frame base64 hop le", "Pass"],
                        ["Validation", "Schema va response code", "Pass"],
                        ["Analyze", "Goi Vision mock va tao analytics_event", "Pass"],
                        ["Newman", "8 requests, 15 assertions", "0 failed"],
                    ],
                )
            if title == "3.6. Bang doi chieu yeu cau":
                table(
                    doc,
                    ["Yeu cau", "Trang thai", "Bang chung"],
                    [
                        ["REST sync Camera -> Vision", "Da co", "POST /api/v1/frames/{id}/analyze"],
                        ["Anh base64 gui sang Vision", "Da co", "image_base64 trong VisionDetectRequest"],
                        ["Queue/MQTT Camera -> Analytics", "Da demo", "publish_camera_event_demo.py"],
                        ["Docker Compose", "Da co", "API, DB, Vision mock, Analytics mock"],
                        ["Newman/GitHub evidence", "Da co", "newman-lab05-compose.html/xml"],
                        ["AI model that", "Chua", "Dang dung Vision mock trong pham vi lab"],
                    ],
                )
            pb(doc)

    h(doc, "KET LUAN")
    for text in [
        "Bao cao da trinh bay day du vai tro cua Camera Stream Service trong Smart Campus Operations Platform. Service duoc xac dinh la thanh phan thu nhan va xu ly so bo du lieu camera, co nhiem vu chuyen frame co chuyen dong thanh request phan tich cho AI Vision va event thong ke cho Analytics.",
        "Ve mat ky thuat, nhom da co hop dong OpenAPI, API upload/analyze frame, Docker Compose, mock Vision/Analytics, Newman evidence va script demo chup frame live/publish MQTT. Phan bo sung anh base64 giup nhom Vision nhan duoc du lieu truc tiep ma khong can file server dung chung.",
        "He thong hien tai da dap ung muc tieu demo va bao cao, nhung van can phan biet ro dau la mock va dau la phan production can phat trien tiep. Huong phat trien la worker motion lien tuc, retry/durable queue va tich hop Vision that.",
    ]:
        p(doc, text)
    pb(doc)

    h(doc, "TAI LIEU THAM KHAO")
    for item in [
        "Nghiep vu 7 service.pdf - tai lieu mo ta nghiep vu cac nhom trong Smart Campus Operations Platform.",
        "Dependency Map - Smart Campus Operations Platform, mapping consumer/provider/co che tich hop.",
        "contracts/camera-stream.openapi.yaml - OpenAPI contract cua Camera Stream Service.",
        "docs/BUSINESS_ANALYSIS_CAMERA.md - phan tich nghiep vu Camera Stream va luong base64 sang Vision.",
        "reports/newman-lab05-compose.html va reports/newman-lab05-compose.xml - bang chung kiem thu Newman.",
        "scripts/auto_capture_camera.py, scripts/publish_camera_event_demo.py, scripts/auto_capture_camera_mqtt_analytics.py - script demo.",
    ]:
        doc.add_paragraph(item, style="List Number")
    pb(doc)

    h(doc, "PHU LUC A. LENH CHAY DEMO")
    code(
        doc,
        "cd C:\\Projects\\Bailap-dich_Vu_Ket_Noi\\lab-5-minhtan39\n"
        "docker compose up -d --build\n"
        "docker compose ps\n"
        "curl http://localhost:8000/health\n"
        "python scripts\\auto_capture_camera.py\n"
        "python scripts\\auto_capture_camera_mqtt_analytics.py",
    )
    p(doc, "Phu luc nay dung khi thuyet trinh truc tiep. Neu demo voi nhom khac, thay localhost bang IP may demo tren cung mang/hotspot.")
    pb(doc)

    h(doc, "PHU LUC B. PAYLOAD GUI CHO VISION")
    code(
        doc,
        '{\n  "request_id": "vision-FR-20260620-0001",\n  "camera_id": "CAM-A01",\n  "timestamp": "2026-06-20T08:30:00+07:00",\n  "location": "Main lobby",\n  "motion_score": 0.82,\n  "image_base64": "base64-jpeg-content",\n  "snapshot_url": null\n}',
    )
    p(doc, "Payload nay la cau tra loi khi team Vision hoi nhom Camera se gui du lieu gi. image_base64 la noi dung anh, cac field con lai la context va khoa doi soat.")
    pb(doc)

    h(doc, "PHU LUC C. PAYLOAD GUI CHO ANALYTICS")
    code(
        doc,
        '{\n  "event_type": "camera.motion.analyzed",\n  "source_service": "team-camera",\n  "frame_id": "FR-20260620-0001",\n  "camera_id": "CAM-A01",\n  "location": "Main lobby",\n  "motion_score": 0.82,\n  "motion_level": "high",\n  "risk_level": "high",\n  "unknown_person": true,\n  "alert_candidate": true\n}',
    )
    p(doc, "Payload nay moi la du lieu da xu ly gui cho Analytics. Khong gui image_base64 sang Analytics vi Analytics chi can aggregate KPI.")
    pb(doc)

    h(doc, "PHU LUC D. MA TRAN TEST CHI TIET")
    table(
        doc,
        ["ID", "Kich ban", "Du lieu vao", "Ket qua mong doi", "Trang thai"],
        [
            ["TC01", "GET /health", "Khong can token", "status ok", "Pass"],
            ["TC02", "POST frame hop le", "camera_id, image_base64, motion_score", "201 FrameAccepted", "Pass"],
            ["TC03", "POST frame sai format", "frame_format jpg", "422 ValidationError", "Da sua thanh jpeg"],
            ["TC04", "Analyze frame ton tai", "frame_id hop le", "Vision response va analytics_event", "Pass"],
            ["TC05", "Analyze frame khong ton tai", "frame_id sai", "404 NotFound", "Can kiem thu them"],
            ["TC06", "Vision unavailable", "Tat Vision mock", "502/503 Problem Details", "Can kiem thu them"],
            ["TC07", "Publish MQTT Analytics", "Event camera.motion.analyzed", "Subscriber nhan message", "Da demo bang script"],
            ["TC08", "Live camera capture", "Camera source URL", "Tao auto-camera-frame.jpg va upload API", "Da demo"],
        ],
    )
    pb(doc)

    h(doc, "PHU LUC E. GOI Y NOI KHI THUYET TRINH")
    for text in [
        "Nhom em lam Camera Stream Service. Service nay khong lam AI va khong lam Analytics, ma nam giua camera source va cac service phan tich. Khi co chuyen dong, Camera chup frame, gan metadata, ma hoa anh base64 va gui sang Vision.",
        "Day la Docker Compose cua nhom em gom API, database, Vision mock va Analytics mock. Endpoint /health cho thay service dang san sang. Sau khi upload frame, API tra frame_id. Khi goi analyze, service tao request sang Vision va nhan ket qua detect.",
        "Day la JSON em gui cho team Vision. Trong do image_base64 la anh, con camera_id, location, timestamp, motion_score la context. Day la JSON em gui cho team Analytics. No khong chua anh, chi chua event da xu ly de Analytics aggregate KPI.",
        "Bai cua nhom em dap ung luong Camera -> Vision REST sync va Camera -> Analytics async. Hien tai Vision la mock de tich hop lab, huong phat trien la thay bang Vision that va hoan thien worker phat hien chuyen dong lien tuc.",
    ]:
        p(doc, text)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Camera Stream Service - FIT4110")
    doc.save(OUT)
    print(OUT)
    print(OUT.stat().st_size)


if __name__ == "__main__":
    main()
